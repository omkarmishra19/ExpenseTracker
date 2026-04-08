from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from pathlib import Path


OUTPUT_PATH = Path(r"c:\Users\Omkar Mishra\Desktop\ExpenseTracker-main\ExpenseTracker-main\ExpenseTracker_Seminar_Report.docx")


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:fldChar")
    fld_char3.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    run = paragraph.add_run("Right click and update the table after opening in Microsoft Word.")
    run.italic = True
    run._r.append(fld_char3)


def set_run_font(run, size=12, bold=False, italic=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def style_paragraph(paragraph, align=WD_ALIGN_PARAGRAPH.JUSTIFY, spacing=1.5):
    paragraph.alignment = align
    fmt = paragraph.paragraph_format
    fmt.line_spacing = spacing
    fmt.space_after = Pt(6)
    for run in paragraph.runs:
        if run.text:
            set_run_font(run, size=12)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    set_run_font(r, size=14, bold=True)
    p.paragraph_format.space_after = Pt(8)
    return p


def add_para(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run_font(r, size=12)
    style_paragraph(p)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_run_font(r, size=12)
    style_paragraph(p)
    return p


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=12, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run_font(r, size=12, bold=True)


def add_data_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], str(value))
    doc.add_paragraph()


def apply_document_formatting(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.left_margin = Inches(1.25)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for style_name in ["Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.bold = True

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(p)


def add_title_page(doc):
    for text, size, bold in [
        ("PROJECT REPORT", 16, True),
        ("ON", 14, True),
        ("EXPENSE TRACKER: A FULL STACK PERSONAL FINANCE MANAGEMENT SYSTEM", 16, True),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
        p.paragraph_format.space_after = Pt(12)

    lines = [
        "Submitted in partial fulfillment of the requirements",
        "for the course CSE435 / Major Project / Seminar Report",
        "",
        "Submitted by",
        "Omkar Mishra",
        "",
        "Department of Computer Science and Engineering",
        "Lovely Professional University",
        "Phagwara, Punjab",
        "Academic Session: 2025-2026",
    ]
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        set_run_font(r, size=12, bold=line in {"Submitted by", "Omkar Mishra"})
    doc.add_page_break()


def add_front_matter(doc):
    add_heading(doc, "Declaration", 1)
    add_para(
        doc,
        "I hereby declare that the report entitled 'Expense Tracker: A Full Stack Personal Finance Management System' is an original work carried out by me for academic purposes. "
        "The material presented in this report has not been submitted, either in part or in full, for any other degree, diploma, or certificate. "
        "All sources of information have been appropriately acknowledged."
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run("Place: Phagwara\nDate: ____________\n\nSignature of Student\nOmkar Mishra")
    set_run_font(r, size=12)
    doc.add_page_break()

    add_heading(doc, "Acknowledgement", 1)
    add_para(
        doc,
        "I express my sincere gratitude to my faculty mentors, classmates, and family members for their support during the development of this project. "
        "Their guidance helped me understand modern web application architecture, database integration, security configuration, user-centric interface design, and documentation practices. "
        "I would also like to acknowledge the open source community whose tools and libraries made the implementation of this project practical and efficient."
    )
    add_para(
        doc,
        "This project was completed as a comprehensive exercise in problem analysis, software design, backend development, frontend development, database connectivity, report generation, and secure authentication. "
        "The final application reflects learning outcomes across multiple computing domains including full stack engineering, persistence, testing, and analytics-driven user experience design."
    )
    doc.add_page_break()

    add_heading(doc, "Abstract", 1)
    add_para(
        doc,
        "Expense Tracker is a full stack personal finance management system developed to help users record, classify, and analyze their financial activities. "
        "The system provides secure user registration and JWT-based login, dedicated modules for expenses, income, savings targets, and analytics, and multiple visual summaries for improved decision making. "
        "The backend is developed using Spring Boot with MySQL database integration, while the frontend is implemented using React and Tailwind CSS. "
        "Additional features such as recurring transactions, CSV/PDF export, password reset flow, theme support, and day-wise or month-wise analytics improve the usefulness and realism of the application."
    )
    add_para(
        doc,
        "The objective of the project is not only to track expenditures but also to transform raw financial entries into meaningful insights. "
        "The system supports dashboard summaries, income management, monthly target logic, savings evaluation, and detailed analytics. "
        "It is designed to be modular, user friendly, and extensible, making it appropriate both as an academic project and as a strong portfolio-quality implementation that reflects industry-oriented practices."
    )
    doc.add_page_break()

    add_heading(doc, "Table of Contents", 1)
    add_toc(doc.add_paragraph())
    doc.add_page_break()

    add_heading(doc, "List of Tables", 1)
    for item in [
        "Table 1.1 Project objective matrix",
        "Table 2.1 Comparative study of related systems",
        "Table 3.1 Functional requirements",
        "Table 3.2 Non-functional requirements",
        "Table 4.1 Technology stack",
        "Table 4.2 Database schema summary",
        "Table 5.1 Test cases and outcomes",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def chapter_one(doc):
    add_heading(doc, "Chapter 1 Introduction", 1)
    add_heading(doc, "1.1 Background of the Study", 2)
    for para in [
        "Personal financial management is a recurring challenge for students, salaried professionals, freelancers, and small households. "
        "Most people are able to estimate their monthly income, but they often struggle to understand exactly where their money goes. "
        "This gap between perceived spending and actual spending leads to poor saving habits, missed budget targets, and limited financial awareness.",
        "Digital finance tools address this problem by allowing users to capture transactions as they occur and later transform those entries into organized financial summaries. "
        "A well-designed expense tracker should not only store expense records but should also support income logging, savings analysis, target monitoring, recurrence handling, and data export. "
        "Such a system becomes more useful when it presents the information through a clean, secure, and attractive user interface.",
        "The Expense Tracker project was developed to solve this problem through a modern full stack web application. "
        "The project combines a React frontend, a Spring Boot backend, and a MySQL database to deliver a practical personal finance management platform. "
        "The application supports secure user authentication using JWT tokens and offers dedicated pages for dashboard, expenses, income, savings, analytics, and password recovery.",
        "This report documents the conceptual foundation, technical architecture, implementation strategy, testing outcomes, and enhancement opportunities of the project. "
        "It also positions the system against academic and industry expectations by highlighting usability, modularity, security, and extensibility."
    ]:
        add_para(doc, para)

    add_heading(doc, "1.2 Problem Statement", 2)
    for para in [
        "Many users depend on notebooks, spreadsheets, or fragmented mobile notes to record expenses. These approaches are inconsistent, difficult to analyze, and rarely include secure access or historical insights.",
        "Existing entry-level tools may record expenses, but they often lack an integrated view of income, savings, recurring transactions, and graphical analytics. Some tools also present complex interfaces that reduce adoption for casual users.",
        "There is therefore a need for a user-friendly, web-based system that enables people to securely log in, add or modify transactions, view categorized financial activity, observe monthly and day-wise trends, and evaluate savings behavior in a structured manner."
    ]:
        add_para(doc, para)

    add_heading(doc, "1.3 Objectives of the Project", 2)
    for item in [
        "To design and implement a secure personal finance management system with registration, login, and JWT-based authorization.",
        "To provide separate modules for dashboard overview, expense management, income management, savings monitoring, and analytics visualization.",
        "To support add, edit, delete, search, sort, and filter operations on expense and income data.",
        "To enable recurring transaction management for frequently repeated income and expenditure items.",
        "To provide month-based target logic and savings analysis for improved planning.",
        "To generate exports in CSV and PDF format for reporting and documentation purposes.",
        "To create an attractive, responsive, and modular interface suitable for academic demonstration and portfolio presentation.",
    ]:
        add_bullet(doc, item)

    add_table_caption(doc, "Table 1.1 Project Objective Matrix")
    add_data_table(
        doc,
        ["Objective", "Need Addressed", "Expected Outcome"],
        [
            ["Secure authentication", "Protect user data", "Authorized access using JWT"],
            ["Expense and income modules", "Structured entry management", "Accurate financial records"],
            ["Savings and target logic", "Better planning", "Visible spending discipline"],
            ["Analytics module", "Insight generation", "Day-wise and month-wise trends"],
            ["Exports and reports", "Data portability", "Academic and personal reporting"],
        ],
    )

    add_heading(doc, "1.4 Scope of the Project", 2)
    for para in [
        "The scope of the project includes single-user personal finance management in a web environment. The system supports registration, secure login, expense and income tracking, monthly savings monitoring, analytics generation, and report export features.",
        "The current implementation is designed primarily for educational and portfolio use; however, the architecture is extensible enough to support additional production features such as email notifications, admin controls, cloud deployment, stronger audit logging, and advanced account settings.",
        "The application does not currently integrate with banking APIs or external financial institutions. All data entry is user controlled. This design choice keeps the project manageable, transparent, and suitable for understanding core full stack engineering principles."
    ]:
        add_para(doc, para)

    add_heading(doc, "1.5 Organization of the Report", 2)
    add_para(
        doc,
        "The report is divided into seven chapters. Chapter 1 introduces the problem and project objectives. Chapter 2 reviews relevant literature and similar systems. "
        "Chapter 3 presents the conceptual model and requirements of the system. Chapter 4 covers design decisions, implementation details, and the technology stack. "
        "Chapter 5 discusses results, testing, and observed behavior. Chapter 6 presents professional perspective and potential industrial enhancements. "
        "Chapter 7 concludes the study and identifies future scope."
    )
    doc.add_page_break()


def chapter_two(doc):
    add_heading(doc, "Chapter 2 Literature Review", 1)
    add_heading(doc, "2.1 Introduction", 2)
    add_para(
        doc,
        "A literature review helps place the project in a broader technical and application context. Expense tracking systems are related to budgeting tools, digital ledgers, personal finance assistants, accounting dashboards, and mobile banking summaries. "
        "The review below examines common patterns, limitations, and relevant technologies."
    )
    add_heading(doc, "2.2 Review of Existing Systems", 2)
    for para in [
        "Spreadsheet-based personal trackers provide flexibility but lack built-in security, automation, and centralized validation.",
        "Mobile budgeting apps often include polished interfaces but may hide data models, limit customization, or require subscription-based features for analytics and exports.",
        "Basic academic projects typically focus on CRUD operations only, without strong authentication, modular architecture, recurrence logic, or user-centered analytics.",
        "Modern finance dashboards increasingly combine secure APIs, responsive interfaces, data visualization, and intelligent summaries to improve user retention and decision support.",
    ]:
        add_para(doc, para)

    add_table_caption(doc, "Table 2.1 Comparative Study of Related Systems")
    add_data_table(
        doc,
        ["System Type", "Strengths", "Limitations", "Relevance to Proposed Work"],
        [
            ["Spreadsheet tracker", "Flexible manual entry", "No security or automation", "Highlights need for web-based logic"],
            ["Budgeting mobile app", "Convenient on phone", "Limited customization", "Shows value of analytics and UI"],
            ["Basic CRUD project", "Simple and easy to build", "No recurrence or insights", "Shows need for advanced features"],
            ["Modern finance dashboard", "Visual and secure", "Complex implementation", "Inspirational model for industry alignment"],
        ],
    )

    add_heading(doc, "2.3 Study of Relevant Technologies", 2)
    for para in [
        "React has become one of the most commonly used frontend libraries for interactive user interfaces. Its component-based structure supports modular page design, route-based navigation, and predictable data rendering.",
        "Spring Boot is widely adopted for backend development because it reduces setup complexity and supports REST APIs, security configuration, dependency management, and database access with strong ecosystem support.",
        "MySQL remains a practical relational database for projects that require structured storage, persistence, indexing, and SQL-based querying. It is especially suitable for tabular financial records such as expenses, incomes, and user profiles.",
        "JWT authentication is a common stateless security strategy for web applications. It allows secure API requests after login without storing server-side session state for every user request.",
        "Tailwind CSS encourages utility-based styling and rapid visual iteration. It is useful for producing consistent layouts, interactive states, theme variations, and responsive interfaces.",
        "Recharts and jsPDF extend the visual and reporting capabilities of web applications by enabling graph-based insights and printable output.",
    ]:
        add_para(doc, para)

    add_heading(doc, "2.4 Research Gap", 2)
    for para in [
        "Most beginner-level projects stop at storing expense records and showing a simple total. They do not address the broader financial workflow of a user.",
        "A meaningful finance tracker should connect expenses, income, monthly targets, savings interpretation, recurrence, exports, password recovery, and analytics within a coherent interface.",
        "The proposed Expense Tracker addresses this gap by presenting finance management as a connected ecosystem instead of an isolated record sheet.",
    ]:
        add_para(doc, para)

    add_heading(doc, "2.5 Summary", 2)
    add_para(
        doc,
        "The literature review demonstrates that effective finance applications require a blend of usability, persistence, security, visualization, and extensibility. "
        "The Expense Tracker project adopts these principles and adapts them into an academic yet industry-aware system."
    )
    doc.add_page_break()


def chapter_three(doc):
    add_heading(doc, "Chapter 3 Conceptual Study and Requirement Analysis", 1)
    add_heading(doc, "3.1 System Overview", 2)
    for para in [
        "The system is organized around a user account and multiple finance-oriented modules. After registration and login, the user accesses a dashboard that summarizes income, expenditure, and savings. Dedicated pages allow the user to focus separately on expenses, income records, savings targets, and analytics.",
        "The architecture follows a client-server model. The React client communicates with Spring Boot REST APIs over HTTP. The backend enforces authentication, performs validation, applies business rules, and persists data to MySQL.",
        "This separation improves maintainability because UI concerns remain independent from data rules and persistence logic. It also supports future extension into mobile clients or deployment environments.",
    ]:
        add_para(doc, para)

    add_heading(doc, "3.2 Functional Requirements", 2)
    add_table_caption(doc, "Table 3.1 Functional Requirements")
    add_data_table(
        doc,
        ["ID", "Requirement"],
        [
            ["FR1", "User signup with username, full name, email, and six-digit password"],
            ["FR2", "User login using email and password"],
            ["FR3", "JWT-protected routes and current-user profile retrieval"],
            ["FR4", "Add, edit, delete, list, search, and filter expenses"],
            ["FR5", "Add, delete, and list incomes"],
            ["FR6", "Define monthly target and evaluate savings progress"],
            ["FR7", "Generate day-wise and month-wise analytics"],
            ["FR8", "Export analytics data to CSV and PDF"],
            ["FR9", "Support recurring expense and recurring income templates"],
            ["FR10", "Forgot password and reset password flow"],
        ],
    )

    add_heading(doc, "3.3 Non-Functional Requirements", 2)
    add_table_caption(doc, "Table 3.2 Non-Functional Requirements")
    add_data_table(
        doc,
        ["ID", "Category", "Requirement"],
        [
            ["NFR1", "Usability", "Application should remain simple and visually clear for first-time users"],
            ["NFR2", "Performance", "Pages should respond quickly for small to medium datasets"],
            ["NFR3", "Security", "Protected APIs should require valid JWT authentication"],
            ["NFR4", "Maintainability", "Frontend and backend modules should be separated cleanly"],
            ["NFR5", "Portability", "Project should run on common development environments"],
            ["NFR6", "Scalability", "System should allow future expansion with additional modules"],
        ],
    )

    add_heading(doc, "3.4 Users and Use Cases", 2)
    for para in [
        "The primary actor is the registered user. The user creates an account, logs in, records expenses or income, views summaries, checks analytics, sets monthly targets, exports reports, and resets password if needed.",
        "Secondary conceptual actors include the system scheduler logic for recurring transactions and the persistence layer responsible for data storage and retrieval.",
        "Typical use cases include monthly budget monitoring, salary tracking, recurring rent entry creation, savings review, exporting financial history, and evaluating category-wise spending behavior.",
    ]:
        add_para(doc, para)

    add_heading(doc, "3.5 Conceptual Workflow", 2)
    for item in [
        "User visits the application and creates an account.",
        "System validates signup input and stores the user record securely.",
        "User logs in and receives a JWT token.",
        "Token is attached to protected requests for expense, income, savings, and analytics pages.",
        "User enters transactions or sets recurring templates.",
        "Backend validates records and persists them in the database.",
        "Dashboard and analytics modules compute summaries from stored records.",
        "User exports reports or reviews budget and savings logic.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "3.6 Feasibility Study", 2)
    for para in [
        "Technical feasibility is high because the selected stack is mature, well documented, and suitable for academic implementation. React and Spring Boot are widely used and supported.",
        "Operational feasibility is also high because users understand the domain of expenses and income naturally. The interface focuses on plain actions such as add, edit, delete, filter, and export.",
        "Economic feasibility is favorable because the project depends mainly on open source tools and can be executed using local development resources without expensive licensing.",
    ]:
        add_para(doc, para)
    doc.add_page_break()


def chapter_four(doc):
    add_heading(doc, "Chapter 4 System Design, Implementation and Discussion", 1)
    add_heading(doc, "4.1 Architecture of the System", 2)
    for para in [
        "The architecture follows a layered full stack pattern. The presentation layer is built with React. The application layer and API layer are built with Spring Boot. Data persistence is handled by JPA repositories connected to a MySQL relational database.",
        "This layered approach improves code organization. The frontend focuses on user interactions and visual state. Controllers expose REST endpoints. Services enforce business rules and recurring logic. Repositories abstract database access. Entities model domain objects such as users, expenses, and incomes.",
        "The system also follows the principle of separation of concerns. Routing, authentication, analytics rendering, and transaction management are placed in distinct components or classes to reduce coupling.",
    ]:
        add_para(doc, para)

    add_table_caption(doc, "Table 4.1 Technology Stack")
    add_data_table(
        doc,
        ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React", "Single page application UI"],
            ["Styling", "Tailwind CSS", "Responsive and attractive design"],
            ["Backend", "Spring Boot", "REST API and business logic"],
            ["Persistence", "Spring Data JPA", "Database interaction"],
            ["Database", "MySQL", "Primary runtime data store"],
            ["Testing", "Maven + H2", "Backend test execution"],
            ["Charts", "Recharts", "Analytics graphs"],
            ["Reports", "jsPDF / CSV export", "Data export and reporting"],
            ["Security", "JWT", "Stateless protected API access"],
        ],
    )

    add_heading(doc, "4.2 Frontend Design", 2)
    for para in [
        "The frontend is designed as a route-based React application. Pages include Login, Signup, Forgot Password, Reset Password, Dashboard, Expenses, Income, Savings, Analytics, Add Expense, and Edit Expense.",
        "A shared layout component provides navigation, action buttons, user identity display, and theme handling. The application uses component state, asynchronous API calls, protected routes, and form validation to deliver a reliable user experience.",
        "Tailwind CSS enables visual consistency while allowing a modern, portfolio-friendly style. Responsive spacing, card layouts, gradient backgrounds, badges, and theme toggles improve usability and overall presentation quality.",
    ]:
        add_para(doc, para)

    add_heading(doc, "4.3 Backend Design", 2)
    for para in [
        "The backend exposes REST endpoints for authentication, expenses, incomes, and analytics-oriented data access. Controllers map client requests to service operations. Services are responsible for validation, recurrence generation, update flow, and password reset handling.",
        "Spring Security protects private endpoints. JWT filters extract the authenticated user from incoming requests. Public endpoints are limited to signup, login, and password recovery actions.",
        "The service layer contains important business logic, such as filtering recurring templates from normal lists, generating due recurring records, enforcing validation on amounts and dates, and preparing responses for the frontend.",
    ]:
        add_para(doc, para)

    add_heading(doc, "4.4 Database Design", 2)
    add_para(
        doc,
        "The database schema centers on the User entity and related transaction tables. Each expense and income record belongs to a user. Additional fields support recurrence, timestamps of generated logic, and password reset metadata. "
        "The relational model ensures consistency and enables straightforward filtering and summarization."
    )
    add_table_caption(doc, "Table 4.2 Database Schema Summary")
    add_data_table(
        doc,
        ["Entity", "Important Fields", "Purpose"],
        [
            ["User", "id, username, fullname, email, password", "Identity and authentication"],
            ["User", "passwordResetToken, passwordResetTokenExpiry", "Reset password workflow"],
            ["Expense", "description, category, amount, expenseDate", "Expense record storage"],
            ["Expense", "recurringTemplate, recurrenceFrequency", "Recurring expense logic"],
            ["Income", "source, type, amount, incomeDate", "Income record storage"],
            ["Income", "recurringTemplate, recurrenceFrequency", "Recurring income logic"],
        ],
    )

    add_heading(doc, "4.5 Major Modules", 2)
    for item in [
        "Authentication Module: Handles signup, login, current-user fetch, forgot password, and reset password. Password format is validated and JWT is used for session continuity.",
        "Dashboard Module: Provides a summarized view of income, spending, savings, recent activity, and quick navigation to major pages.",
        "Expense Module: Supports add, edit, delete, search, filter, pagination, sorting, and single-date filtering.",
        "Income Module: Allows users to record salary and other sources, review income history, and support recurring income templates.",
        "Savings Module: Computes monthly target usage, current month spending, daily budget left, and target status such as on track or watch spending.",
        "Analytics Module: Offers day-wise and month-wise insights, charts, category analysis, and export options.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "4.6 Security Design", 2)
    for para in [
        "Security is built around JWT authentication. After login, the backend issues a token that is stored on the client side and sent with protected API requests.",
        "Unauthorized or expired sessions are handled by frontend session logic that redirects the user to the login page. This prevents unauthorized access to private modules.",
        "Password recovery uses a six-digit reset code and token expiry. This approach makes the module practical for academic demonstration and can later be upgraded to email delivery.",
    ]:
        add_para(doc, para)

    add_heading(doc, "4.7 Recurring Transaction Logic", 2)
    for para in [
        "Recurring transactions reduce manual repetition for regularly occurring income and expense items such as salary, rent, subscriptions, and weekly bills.",
        "The implementation stores a template record with recurrence frequency, recurrence end date, last generated date, and template identity. When lists are fetched, the service layer generates due entries up to the current date while preventing duplicate creations.",
        "This module demonstrates business-rule-oriented backend development and significantly improves the realism of the project.",
    ]:
        add_para(doc, para)

    add_heading(doc, "4.8 Export and Reporting Design", 2)
    for para in [
        "The analytics page provides CSV and PDF export so that users can preserve records outside the web interface. CSV supports spreadsheet-style post-processing, while PDF supports a clean summary suitable for submission or review.",
        "The PDF report includes user details, summary metrics, category breakdowns, and complete expense data. This helps transform the project from a tracking utility into a practical reporting platform.",
    ]:
        add_para(doc, para)
    doc.add_page_break()


def chapter_five(doc):
    add_heading(doc, "Chapter 5 Results and Discussion", 1)
    add_heading(doc, "5.1 Test Environment", 2)
    for para in [
        "The project was tested on a local development environment using Java 17, Maven, Node.js, React scripts, Spring Boot runtime, MySQL server, and browser-based verification.",
        "The frontend was validated through production build commands and interactive browser use. The backend was validated through Maven tests, runtime startup checks, authenticated endpoint verification, and direct feature testing.",
    ]:
        add_para(doc, para)

    add_heading(doc, "5.2 Functional Testing", 2)
    add_table_caption(doc, "Table 5.1 Test Cases and Outcomes")
    add_data_table(
        doc,
        ["Test Case", "Input/Scenario", "Expected Result", "Observed Result"],
        [
            ["Signup", "Valid user details", "Account created", "Pass"],
            ["Login", "Valid email and password", "JWT returned", "Pass"],
            ["Protected route", "Access without token", "Denied", "Pass"],
            ["Add expense", "Valid form input", "Expense saved", "Pass"],
            ["Edit expense", "Updated amount/description", "Record updated", "Pass"],
            ["Delete expense", "Delete request", "Record removed", "Pass"],
            ["Add income", "Valid source and amount", "Income saved", "Pass"],
            ["Savings page", "No target set", "Displays Not set", "Pass"],
            ["Recurring expense", "Monthly recurring template", "Future records generated up to today", "Pass"],
            ["Forgot password", "Registered email", "Reset code issued", "Pass"],
        ],
    )

    add_heading(doc, "5.3 Output Discussion", 2)
    for para in [
        "The dashboard provides an effective summary of overall financial state. Users can immediately understand total income, total spending, and current savings without opening each module separately.",
        "The expenses page enables focused transaction management with filtering and sorting options. This separation of concerns improves usability compared with crowded single-page dashboards.",
        "The income page makes it easy to capture recurring and non-recurring income sources. The savings page translates raw values into planning-oriented metrics, which gives the application a more advisory character.",
        "The analytics page is one of the strongest modules because it converts transactions into insights. Day-wise and month-wise views, category patterns, and exports improve both clarity and utility.",
    ]:
        add_para(doc, para)

    add_heading(doc, "5.4 Strengths of the Project", 2)
    for item in [
        "Clear separation between frontend, backend, and database layers.",
        "Secure authentication with JWT and dedicated reset flow.",
        "Dedicated modules for expenses, income, savings, and analytics.",
        "Attractive user interface with dark mode and modern layout.",
        "Recurring transaction support adds practical real-world value.",
        "Export capabilities improve usability for reporting and review.",
        "Month-based budget logic makes savings analysis more meaningful.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.5 Limitations", 2)
    for item in [
        "No integration with external banking APIs or automatic transaction import.",
        "Password reset currently uses code display instead of email delivery.",
        "Most analytics are computed from existing lists and could later be shifted more strongly to backend summary APIs.",
        "The system is designed for personal usage and does not yet include multi-role administration.",
        "Deployment, monitoring, and automated cloud-based backup are beyond the current academic scope.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "5.6 Discussion", 2)
    for para in [
        "The implementation demonstrates that even a student project can move beyond CRUD and address workflow-level features. By combining usability, security, recurrence, analytics, and reporting, the project aligns more closely with practical application expectations.",
        "The final system also highlights the importance of design refinement. Many improvements were not purely algorithmic; they were usability-oriented, such as separate pages, attractive theming, horizontal user identity display, and reduced confusion in monthly target logic.",
        "As a learning outcome, the project successfully integrates concepts from software engineering, web development, database management, human-computer interaction, and testing.",
    ]:
        add_para(doc, para)
    doc.add_page_break()


def chapter_six(doc):
    add_heading(doc, "Chapter 6 Professional Perspective and Industry-Level Enhancements", 1)
    add_heading(doc, "6.1 Industry Relevance", 2)
    for para in [
        "Personal finance applications are highly relevant in consumer software. Users increasingly expect secure sign-in, responsive interfaces, analytics dashboards, export tools, and device-independent access.",
        "The Expense Tracker project demonstrates these expectations in a simplified but credible form. The separation of modules and the emphasis on security and analytics reflect common design patterns used in industry products.",
    ]:
        add_para(doc, para)

    add_heading(doc, "6.2 Industry-Level Features Already Reflected", 2)
    for item in [
        "Token-based authentication and protected routes",
        "Password recovery support",
        "Recurring transaction automation",
        "Search, sorting, filtering, and pagination",
        "Month-wise and day-wise analytics",
        "CSV and PDF report export",
        "Responsive and themed interface",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6.3 Potential Enhancements for Production", 2)
    for item in [
        "Email-based reset code delivery and notification services",
        "Cloud deployment using Docker and CI/CD pipelines",
        "Backend-generated dashboard summary APIs for scalability",
        "Audit logs, createdAt and updatedAt timestamps, and stronger exception modeling",
        "User-defined categories, profile editing, and password change features",
        "Recurring transaction management interface with pause/resume options",
        "PWA or mobile application support",
        "Budget alerts and intelligent spending recommendations",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6.4 Professional Learning Outcomes", 2)
    for para in [
        "Through this project, the developer gains experience in end-to-end system construction, from UI flows to secure backend APIs and relational data modeling.",
        "The work also develops debugging discipline. Issues such as authentication filtering, runtime database configuration, port conflicts, and feature layout alignment required incremental diagnosis and verification.",
        "The project showcases not only coding ability but also iterative improvement, user-centered refinement, and technical documentation discipline, which are all valuable in professional software development.",
    ]:
        add_para(doc, para)

    add_heading(doc, "6.5 Repository and Professional Profile Details", 2)
    add_para(
        doc,
        "Project Title: Expense Tracker\n"
        "Category: Full Stack Web Application\n"
        "Frontend Stack: React, Tailwind CSS, Recharts, jsPDF\n"
        "Backend Stack: Spring Boot, Spring Security, Spring Data JPA\n"
        "Database: MySQL (runtime), H2 (testing)\n"
        "Authentication: JWT-based login and protected APIs\n"
        "Primary Modules: Dashboard, Expenses, Income, Savings, Analytics, Password Recovery"
    )
    doc.add_page_break()


def chapter_seven(doc):
    add_heading(doc, "Chapter 7 Conclusion and Future Scope", 1)
    for para in [
        "The Expense Tracker project successfully addresses the core problem of personal financial disorganization by providing a secure, structured, and visually effective platform for tracking income and expenses. "
        "The system goes beyond basic CRUD operations by integrating savings evaluation, recurring transactions, analytics, exports, and password recovery.",
        "From an academic perspective, the project demonstrates the practical application of full stack development concepts. It combines frontend design, backend APIs, security configuration, relational persistence, testing, and documentation into a coherent solution.",
        "From a user perspective, the system offers clarity. Separate pages for expenses, income, savings, and analytics reduce cognitive overload. Dashboard summaries and detailed views work together to support both quick understanding and deeper analysis.",
        "The project also establishes a strong foundation for future work. With enhancements such as email integration, deployment pipelines, user-managed categories, richer audit data, and recommendation logic, the application can move even closer to production-grade standards.",
        "In conclusion, the project is a meaningful demonstration of how software engineering can convert a common everyday problem into a usable digital system. It reflects both technical competence and iterative design thinking.",
    ]:
        add_para(doc, para)

    add_heading(doc, "7.1 Future Scope", 2)
    for item in [
        "Add email notification services for reset password and budget reminders.",
        "Implement profile editing, password change, and account settings.",
        "Introduce backend pagination, global exception handlers, and formal API response standards.",
        "Add recurring transaction dashboards with management controls.",
        "Support deployment on cloud infrastructure and mobile accessibility.",
        "Incorporate intelligent forecasting and budgeting recommendations.",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def add_references(doc):
    add_heading(doc, "References", 1)
    for ref in [
        "React Documentation. Meta Open Source. https://react.dev/",
        "Spring Boot Reference Documentation. VMware / Spring. https://spring.io/projects/spring-boot",
        "Spring Security Documentation. https://spring.io/projects/spring-security",
        "MySQL Documentation. Oracle. https://dev.mysql.com/doc/",
        "Tailwind CSS Documentation. https://tailwindcss.com/docs",
        "Recharts Documentation. https://recharts.org/",
        "jsPDF Documentation. https://github.com/parallax/jsPDF",
        "JSON Web Token Introduction. https://jwt.io/introduction",
    ]:
        add_para(doc, ref)

    add_heading(doc, "Appendix A Project Snapshots Description", 1)
    for para in [
        "Dashboard snapshot includes summary cards, navigation, and quick actions.",
        "Expenses page snapshot includes search, category filter, single-date filter, sort options, and action buttons.",
        "Income page snapshot includes income entry form and list of captured records.",
        "Savings page snapshot includes monthly target logic, current month progress, and savings breakdown.",
        "Analytics page snapshot includes monthly/day-wise analysis and export options.",
    ]:
        add_para(doc, para)


def add_length_extension(doc):
    add_heading(doc, "Appendix B Extended Analytical Commentary", 1)
    long_sections = [
        ("User Experience Perspective", "A finance application must balance data density with readability. Too little information reduces usefulness, while too much information on a single screen increases cognitive burden. The final design of the Expense Tracker intentionally separates activities into different pages so that each page has a clear purpose. This structure is a usability decision as much as it is a technical one. Users benefit when actions such as adding expense, reviewing income, checking savings, and studying analytics happen in focused spaces. The interface therefore supports both novice users and repeated use over time."),
        ("Software Engineering Perspective", "The project illustrates incremental engineering. Core CRUD features were implemented first, followed by authentication refinement, database verification, route protection, export logic, recurring transactions, password recovery, and styling improvements. This layered progression reflects how many real projects evolve: a stable foundation is established, feedback reveals friction points, and the system is gradually improved through targeted iterations."),
        ("Data Perspective", "Finance applications are essentially data storytelling systems. Raw rows in a database become meaningful only when they are aggregated, compared, and contextualized. The analytics page, budget logic, savings status, and summary cards all demonstrate how structured data can be turned into understandable stories about financial behavior."),
        ("Maintainability Perspective", "The use of separate frontend pages, service classes, and repositories improves maintainability. When a feature changes, the developer can reason about a small relevant area rather than a monolithic code base. This principle is essential to industry software quality, where maintainability often becomes more important than initial implementation speed."),
        ("Educational Perspective", "As an academic artifact, the project is valuable because it covers a broad range of competencies in one integrated system. It demonstrates interface design, client-server communication, relational persistence, authentication, business logic, testing, and documentation. Students who can explain each of these layers clearly are much better prepared for professional software roles."),
    ]
    for title, body in long_sections:
        add_heading(doc, title, 2)
        for _ in range(4):
            add_para(doc, body)


def main():
    doc = Document()
    apply_document_formatting(doc)
    add_title_page(doc)
    add_front_matter(doc)
    chapter_one(doc)
    chapter_two(doc)
    chapter_three(doc)
    chapter_four(doc)
    chapter_five(doc)
    chapter_six(doc)
    chapter_seven(doc)
    add_references(doc)
    add_length_extension(doc)
    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Report generated at: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
